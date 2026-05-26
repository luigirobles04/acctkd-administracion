#!/usr/bin/env python3
"""Genera el documento de entrega: diagramas de arquitectura (negocio, aplicaciones, datos)."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
IMG = ROOT / 'documentacion' / 'diagramas_png'
OUT_REPO = ROOT / 'documentacion' / 'DIAGRAMAS_ARQUITECTURA_PRACTICAS_ROBLES_PALACIOS.docx'
OUT_DL = Path.home() / 'Downloads' / 'DIAGRAMAS_ARQUITECTURA_PRACTICAS_ROBLES_PALACIOS.docx'


def setup_doc():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(3)
        s.right_margin = Cm(2.5)
    return doc


def run(doc, text, bold=False, italic=False, size=11):
    r = doc.add_paragraph().add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.paragraphs[-1].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return doc.paragraphs[-1]


def center(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.name = 'Times New Roman'
    return p


def figura(doc, caption, img_path, width_cm=15.5):
    center(doc, caption, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(img_path), width=Cm(width_cm))
    center(doc, 'Fuente: Elaboración propia.', italic=True, size=10)
    doc.add_paragraph()


def main():
    doc = setup_doc()

    # --- Portada ---
    center(doc, 'UNIVERSIDAD CÉSAR VALLEJO', bold=True, size=14)
    center(doc, 'FACULTAD DE INGENIERÍA Y ARQUITECTURA', bold=True, size=12)
    center(doc, 'ESCUELA PROFESIONAL DE INGENIERÍA DE SISTEMAS', bold=True, size=12)
    doc.add_paragraph()
    center(doc, 'PRÁCTICAS PRE-PROFESIONALES I', bold=True, size=13)
    doc.add_paragraph()
    center(doc, 'DIAGRAMAS DE ARQUITECTURA DEL PROYECTO', bold=True, size=14)
    center(doc, 'Sistema Web ERP «ACCTKD» — Christopher Cabrera Taekwondo', bold=True, size=12)
    doc.add_paragraph()
    doc.add_paragraph()

    center(doc, 'Presentado por:', bold=True)
    center(doc, 'Robles Palacios, Luigi Armando')
    center(doc, 'Código / Estudiante — Ingeniería de Sistemas (2026)')
    doc.add_paragraph()

    center(doc, 'Empresa formadora:', bold=True)
    center(doc, 'Christopher Cabrera Taekwondo (ACCTKD)')
    center(doc, 'Jefe inmediato: Christopher Cabrera Nole — Gerente General, 6° Dan')
    center(doc, 'Calle Puerto Rico 302, Urb. El Recreo, Trujillo — La Libertad')
    doc.add_paragraph()

    center(doc, 'Docente de la asignatura:', bold=True)
    center(doc, 'Fernando Santiago Gonzales Zavaleta')
    doc.add_paragraph()
    doc.add_paragraph()

    center(doc, 'Trujillo — Perú, mayo de 2026')

    doc.add_page_break()

    # --- Introducción ---
    heading(doc, 'INTRODUCCIÓN', 1)
    run(
        doc,
        'El presente documento desarrolla los diagramas de arquitectura solicitados para la evaluación '
        'de las prácticas pre-profesionales: arquitectura de negocio, arquitectura de aplicaciones y '
        'arquitectura de datos del proyecto «ACCTKD-Administración», sistema web ERP diseado para '
        'Christopher Cabrera Taekwondo, academia ubicada en Trujillo.',
    )
    run(
        doc,
        'Durante las semanas de práctica identifiqué que la operación diaria de la academia — matrículas, '
        'control de asistencia en nueve turnos semanales y seguimiento de mensualidades — dependía de '
        'cuadernos, hojas sueltas y archivos Excel que no conversaban entre sí. Por ello, como practicante '
        'de Ingeniería de Sistemas, propuse e implementé una solución digital que no solo registra datos, '
        'sino que refleja la forma real en que la organización trabaja. Estos diagramas documentan esa '
        'comprensión en tres niveles complementarios, siguiendo un enfoque alineado a las capas TOGAF: '
        'qué hace el negocio, cómo está construido el software y dónde viven los datos.',
    )
    run(
        doc,
        'La versión 1.0 del sistema, desplegada en mayo de 2026, administra una sede, nueve turnos, '
        'sesenta alumnos, tres planes de mensualidad (S/ 100, S/ 130 y S/ 180), un maestro titular '
        '6° Dan y módulos de login, dashboard, alumnos, asistencia, pagos y maestros. Las figuras que '
        'siguen fueron elaboradas a partir del informe de prácticas, del repositorio '
        'github.com/luigirobles04/acctkd-administracion y de la base productiva en Supabase Cloud.',
    )

    # --- DIAGRAMA 1: NEGOCIO ---
    heading(doc, 'DIAGRAMA 1 — ARQUITECTURA DE NEGOCIO', 1)
    run(
        doc,
        'La arquitectura de negocio responde a la pregunta «¿qué hace ACCTKD y quién interviene?». '
        'No describe servidores ni lenguajes de programación; describe procesos, actores y flujos de '
        'información que el ERP digitaliza. Christopher Cabrera Taekwondo es una microempresa deportiva '
        'cuyo producto central es la formación marcial con valores de disciplina y respeto. Mi trabajo '
        'consistió en mapear cuatro macroprocesos que sostienen esa operación: académico-operativo, '
        'cobranza, seguimiento deportivo y soporte a la gerencia.',
    )

    heading(doc, 'Actores y roles', 2)
    run(
        doc,
        'El gerente y representante legal, Christopher Cabrera Nole (6° Dan), es el usuario principal '
        'del sistema: consulta indicadores, revisa pagos y supervisa alumnos activos. El maestro titular '
        'imparte clases y registra asistencia desde el celular al finalizar cada turno. Yo, como practicante '
        'de sistemas, desarrollo, despliego y mantengo la plataforma bajo su supervisión. Los apoderados '
        'y alumnos son beneficiarios del servicio; en la versión 1.0 la comunicación de cobro se realiza '
        'por WhatsApp, aunque el roadmap prevé un portal para familias en versiones posteriores.',
    )

    heading(doc, 'Capacidades y macroprocesos', 2)
    run(
        doc,
        'Las capacidades de negocio se encadenan de forma natural: la formación marcial exige gestionar '
        'alumnos (matrícula, ficha, plan y turno); eso habilita el control operativo de horarios y '
        'asistencia; a su vez, cada alumno matriculado genera obligaciones financieras que deben '
        'controlarse y reportarse a gerencia. El flujo típico inicia cuando un apoderado muestra interés, '
        'continúa con la matrícula y asignación de plan-turno, sigue con la planificación de nueve horarios '
        'semanales, la impartición de clase y el registro de asistencia (Presente, Ausente, Justificado o '
        'Retirado). En paralelo se emiten cuotas de mensualidad, se registran pagos por Yape, efectivo o '
        'Plin, y el sistema alerta cuando un pago está vencido o próximo a vencer en los siguientes siete '
        'días — regla de negocio que implementé porque el gerente lo solicitó explícitamente.',
    )
    run(
        doc,
        'La Figura 1 resume este recorrido en un mapa de macroprocesos. La Figura 2 ubica al ERP dentro '
        'de la organización: gerencia, dirección académica y área administrativa-ti se relacionan con '
        'apoderados y, en el futuro, con la FDNT para campeonatos.',
    )

    figura(
        doc,
        'Figura 1: Arquitectura de negocio — macroprocesos de la academia (matrícula, operación, cobranza y gestión)',
        IMG / 'fig6_negocio_procesos.png',
        16,
    )
    figura(
        doc,
        'Figura 2: Arquitectura de negocio — contexto organizacional, actores externos e interacción con el ERP',
        IMG / 'fig6b_negocio_contexto.png',
        14,
    )

    # --- DIAGRAMA 2: APLICACIONES ---
    heading(doc, 'DIAGRAMA 2 — ARQUITECTURA DE APLICACIONES', 1)
    run(
        doc,
        'Si la capa de negocio explica el «qué», la arquitectura de aplicaciones explica el «cómo» '
        'está organizado el software. ACCTKD v1.0 adopta una arquitectura JAMStack serverless de tres '
        'capas, seleccionada porque la academia no cuenta con infraestructura propia ni presupuesto para '
        'servidores dedicados: el despliegue en Vercel y la base en Supabase Cloud permiten escalar con '
        'costo casi nulo en etapa inicial y actualizar el sistema con un simple push al repositorio GitHub.',
    )

    heading(doc, 'Capas lógicas y módulos', 2)
    run(
        doc,
        'En la capa de presentación, Next.js 16.2.4 con React 19 y Turbopack renderiza la interfaz '
        'administrativa con un diseño inspirado en iOS (globals.css y Tailwind CSS v4). Los módulos '
        'funcionales son login, dashboard, alumnos, asistencia, pagos y maestros; cada pantalla consume '
        'servicios JavaScript que concentran la lógica de negocio: asistencia.service.js para matrices '
        'multi-turno y modos Día/Semana/Mes/Rango; pagoAlerts.service.js para badges de mensualidad '
        'vencida o pendiente; y servicios de alumno, maestro y autenticación. La capa de integración '
        'utiliza el cliente @supabase/supabase-js, que comunica con PostgREST, autenticación y políticas '
        'Row Level Security en PostgreSQL 17.6. La integración con WhatsApp no usa API oficial: el módulo '
        'de pagos genera enlaces wa.me con el texto del recordatorio, suficiente para la operación actual.',
    )

    heading(doc, 'Despliegue y entorno productivo', 2)
    run(
        doc,
        'El flujo de despliegue es continuo: desarrollo local, commit a la rama main del repositorio '
        'acctkd-administracion, build automático en Vercel (región gru1) y consumo HTTPS desde Trujillo '
        '(UTC-5) por gerente y maestro. La Figura 3 muestra las capas y módulos; la Figura 4 detalla '
        'la cadena GitHub → Vercel → Supabase.',
    )

    figura(
        doc,
        'Figura 3: Arquitectura de aplicaciones — capas lógicas, servicios y módulos funcionales',
        IMG / 'fig7_aplicaciones.png',
        16,
    )
    figura(
        doc,
        'Figura 4: Arquitectura de aplicaciones — vista de despliegue (desarrollo, GitHub, Vercel y Supabase Cloud)',
        IMG / 'fig7b_despliegue.png',
        14,
    )

    # --- DIAGRAMA 3: DATOS ---
    heading(doc, 'DIAGRAMA 3 — ARQUITECTURA DE DATOS', 1)
    run(
        doc,
        'La arquitectura de datos define qué información persiste, cómo se relaciona y quién puede '
        'consultarla. Toda la operación de ACCTKD converge en PostgreSQL 17 administrado por Supabase, '
        'organizado en dominios que reflejan la realidad de la academia: organización (sede, turno, '
        'maestro), académico (alumno, grado, clase, asistencia), financiero (plan, pago, concepto, '
        'método) y seguridad (usuario, rol, RLS). Diseñar el modelo relacional fue uno de los entregables '
        'más exigentes de la práctica, porque debía soportar nueve turnos con titularidad de maestro, '
        'historial de grados y cobros con descuentos sin duplicar montos.',
    )

    heading(doc, 'Modelo entidad-relación y dominios', 2)
    run(
        doc,
        'Las relaciones principales siguen la lógica operativa: una sede tiene varios turnos; cada turno '
        'genera clases por fecha y concentra alumnos; un maestro puede ser titular de uno o más turnos '
        'mediante la tabla puente maestro_turno. Cada alumno contrata un plan de mensualidad, mantiene un '
        'grado marcial actual con historial de promociones, genera pagos y deja marcas de asistencia '
        'vinculadas a clases concretas. El campo monto_final de la tabla pago es una columna GENERADA '
        'en base de datos, lo que evita inconsistencias entre monto, descuento y total cobrado.',
    )
    run(
        doc,
        'La Figura 5 presenta el núcleo del modelo ER. Las Figuras 6 y 7 amplían los dominios académico '
        'y financiero-seguridad, respectivamente, porque en la evaluación resulta más claro separar '
        'asistencia y grados de cobranza y control de acceso por rol.',
    )

    heading(doc, 'Volúmenes y flujo de información', 2)
    run(
        doc,
        'Al cierre de la versión 1.0 (mayo 2026), la base productiva registra: 1 sede, 9 turnos, 60 alumnos '
        '(55 activos), 366 clases con backfill histórico, 770 asistencias, 144 pagos, 15 grados marciales '
        '(incluidos Dan 4° a 9°) y un maestro titular 6° Dan. Los formularios de la interfaz aplican reglas '
        'antes de persistir: el plan sugiere monto, el vencimiento activa badges visuales y el turno '
        'determina qué clases aparecen en la matriz de asistencia. Las consultas alimentan el dashboard, '
        'el resumen de cobranza comparado con el mes anterior (siempre respecto al calendario actual, '
        'no al mes seleccionado en pantalla) y las matrices de asistencia por rango de fechas.',
    )

    figura(
        doc,
        'Figura 5: Arquitectura de datos — modelo entidad-relación resumido (núcleo v1.0)',
        IMG / 'fig8_datos_er.png',
        16,
    )
    figura(
        doc,
        'Figura 6: Arquitectura de datos — dominio académico (alumno, grado, clase y asistencia)',
        IMG / 'fig81_academico.png',
        15,
    )
    figura(
        doc,
        'Figura 7: Arquitectura de datos — dominio financiero y seguridad (pagos, planes y rol de usuario)',
        IMG / 'fig82_cobranza.png',
        15,
    )

    # --- Conclusión ---
    heading(doc, 'CONCLUSIONES', 1)
    run(
        doc,
        'Los tres diagramas presentados — negocio, aplicaciones y datos — describen de forma integrada '
        'el proyecto de prácticas realizado en ACCTKD. La capa de negocio validó que el ERP responde a '
        'procesos concretos de la academia y no a un catálogo genérico de funciones; la capa de aplicaciones '
        'demuestra una solución moderna, accesible desde navegador móvil y mantenible por un practicante '
        'sin equipo de infraestructura; la capa de datos garantiza trazabilidad académica y financiera con '
        'reglas en aplicación y en base de datos.',
    )
    run(
        doc,
        'Este documento cumple el indicador de evaluación al desarrollar las arquitecturas solicitadas '
        'con figuras elaboradas por el practicante, sustentadas en la implementación desplegada en '
        'https://acctkd-administracion.vercel.app y documentadas en el informe PPP1-C1-PRA06a de prácticas '
        'pre-profesionales I.',
    )

    doc.add_paragraph()
    center(doc, 'Robles Palacios, Luigi Armando', bold=True)
    center(doc, 'Practicante — Ingeniería de Sistemas, UCV')
    center(doc, 'Trujillo, mayo de 2026')

    doc.save(OUT_REPO)
    doc.save(OUT_DL)
    print('OK', OUT_REPO)
    print('OK', OUT_DL)


if __name__ == '__main__':
    main()
