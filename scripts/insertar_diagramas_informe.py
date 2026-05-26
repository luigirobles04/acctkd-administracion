#!/usr/bin/env python3
"""Inserta diagramas PNG en el informe de prácticas mejorado."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

SRC = Path('/Users/luigiarmandoroblespalacios/Downloads/PPP1-C1-PRA06a-INFORME_PRACTICAS-ROBLES_PALACIOS_LUIGI_MEJORADO (1).docx')
OUT_REPO = Path(__file__).resolve().parents[1] / 'documentacion' / 'PPP1-C1-PRA06a-INFORME_PRACTICAS-ROBLES_PALACIOS_LUIGI_CON_DIAGRAMAS.docx'
OUT_DL = Path.home() / 'Downloads' / 'PPP1-C1-PRA06a-INFORME_PRACTICAS-ROBLES_PALACIOS_LUIGI_CON_DIAGRAMAS.docx'
IMG = Path(__file__).resolve().parents[1] / 'documentacion' / 'diagramas_png'


def insert_paragraph_before(paragraph):
    new_p = OxmlElement('w:p')
    paragraph._p.addprevious(new_p)
    return Paragraph(new_p, paragraph._parent)


def insert_paragraph_after(paragraph):
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def add_centered_image_before(anchor, img_path, width_cm=15.5):
    p = insert_paragraph_before(anchor)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(img_path), width=Cm(width_cm))
    return p


def add_centered_image_after(anchor, img_path, width_cm=15.5):
    p = insert_paragraph_after(anchor)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(img_path), width=Cm(width_cm))
    return p


def add_text_before(anchor, text, bold=False):
    p = insert_paragraph_before(anchor)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(11)
    return p


def add_text_after(anchor, text, bold=False):
    p = insert_paragraph_after(anchor)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(11)
    return p


def find_paragraph(doc, substr):
    for p in doc.paragraphs:
        if substr in (p.text or ''):
            return p
    return None


def replace_in_paragraphs(doc, replacements):
    for p in doc.paragraphs:
        if not p.text:
            continue
        for old, new in replacements:
            if old in p.text:
                p.text = p.text.replace(old, new)


def main():
    doc = Document(str(SRC))

    # Renumerar secciones 3.2.3+ para abrir espacio a arquitectura de negocio
    replace_in_paragraphs(
        doc,
        [
            ('3.2.7. Cronograma resumido', '3.2.8. Cronograma resumido'),
            ('3.2.6. Pantallas principales', '3.2.7. Pantallas principales'),
            ('3.2.5. Stack tecnológico', '3.2.6. Stack tecnológico'),
            ('3.2.4. Modelo entidad-relación (resumen)', '3.2.5. Arquitectura de datos'),
            ('3.2.3. Arquitectura del sistema', '3.2.4. Arquitectura de aplicaciones'),
        ],
    )

    p_arch = find_paragraph(doc, '3.2.4. Arquitectura de aplicaciones')
    if p_arch:
        add_text_before(
            p_arch,
            '3.2.3. Arquitectura de negocio\n'
            'La arquitectura de negocio describe los macroprocesos de ACCTKD (matrícula, operación de clases, '
            'cobranza y gestión gerencial) y el contexto organizacional con sus actores: gerencia, maestro titular, '
            'practicante de sistemas y apoderados.',
        )
        add_text_before(p_arch, 'Figura 6.1: Arquitectura de negocio — macroprocesos de la academia')
        add_centered_image_before(p_arch, IMG / 'fig6_negocio_procesos.png', 16)
        add_text_before(p_arch, 'Fuente: Elaboración propia.')
        add_text_before(p_arch, 'Figura 6.2: Arquitectura de negocio — contexto organizacional y actores')
        add_centered_image_before(p_arch, IMG / 'fig6b_negocio_contexto.png', 14)
        add_text_before(p_arch, 'Fuente: Elaboración propia.')

    p7 = find_paragraph(doc, 'Figura 7: Arquitectura lógica del sistema')
    if not p7:
        p7 = find_paragraph(doc, 'Figura 7: Arquitectura de aplicaciones')
    if p7:
        p7.text = 'Figura 7: Arquitectura de aplicaciones — capas lógicas, módulos y servicios'
        img7 = add_centered_image_after(p7, IMG / 'fig7_aplicaciones.png', 16)
        p71 = add_text_after(img7, 'Figura 7.1: Arquitectura de aplicaciones — despliegue (GitHub, Vercel, Supabase)')
        add_centered_image_after(p71, IMG / 'fig7b_despliegue.png', 14)

    p8 = find_paragraph(doc, 'Figura 8: Modelo entidad-relación resumido')
    if p8:
        p8.text = 'Figura 8: Arquitectura de datos — modelo entidad-relación (núcleo v1.0)'
        add_centered_image_after(p8, IMG / 'fig8_datos_er.png', 16)

    p81 = find_paragraph(doc, 'Figura 8.1: Modelo académico y asistencia')
    if p81:
        p81.text = 'Figura 8.1: Arquitectura de datos — dominio académico y asistencia'
        add_centered_image_after(p81, IMG / 'fig81_academico.png', 15)

    p82 = find_paragraph(doc, 'Figura 8.2: Modelo de cobranza y seguridad')
    if p82:
        p82.text = 'Figura 8.2: Arquitectura de datos — dominio financiero y seguridad (rol/usuario)'
        add_centered_image_after(p82, IMG / 'fig82_cobranza.png', 15)

    doc.save(str(OUT_REPO))
    doc.save(str(OUT_DL))
    print('OK', OUT_REPO)
    print('OK', OUT_DL)


if __name__ == '__main__':
    main()
