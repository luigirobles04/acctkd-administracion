#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from pathlib import Path

OUT = Path(__file__).resolve().parents[1] / 'documentacion' / 'DIAGRAMAS_ARQUITECTURA_ACCTKD_ENTREGA.docx'
OUT2 = Path.home() / 'Downloads' / 'DIAGRAMAS_ARQUITECTURA_ACCTKD_ROBLES_PALACIOS.docx'

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(3)
    s.right_margin = Cm(2.5)

def H(t, lvl=1):
    p = doc.add_heading(t, level=lvl)
    for r in p.runs:
        r.font.name = 'Times New Roman'

def P(t, bold=False, center=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(t)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    r.bold = bold
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

def mono(t):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.font.name = 'Courier New'
    r.font.size = Pt(9)

def bullet(items):
    for i in items:
        p = doc.add_paragraph(i, style='List Bullet')
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(11)

H('DIAGRAMAS DE ARQUITECTURA', 1)
P('FACULTAD DE INGENIERÍA Y ARQUITECTURA — ESCUELA PROFESIONAL DE INGENIERÍA DE SISTEMAS', center=True)
P('Proyecto: Sistema Web ERP «ACCTKD» — Christopher Cabrera Taekwondo', center=True)
P('Practicante: Robles Palacios, Luigi Armando | Docente: Fernando Santiago Gonzales Zavaleta', center=True)
P('Mayo 2026 — Trujillo, Perú', center=True)
doc.add_paragraph()
P('Desarrollo de diagramas de arquitectura de negocio, arquitectura de aplicaciones y arquitectura de datos del proyecto de prácticas, alineado al informe PPP1-C1-PRA06a y a la rúbrica de evaluación.')

# --- DIAGRAMA 1 ---
H('DIAGRAMA 1 — ARQUITECTURA DE NEGOCIO', 2)
P('Modela procesos, actores y capacidades de la academia sin detallar tecnología.')
H('Actores', 3)
bullet([
    'Gerente: Christopher Cabrera Nole, 6° Dan — reportes, cobranza, supervisión.',
    'Maestro titular: clases y asistencia en 9 turnos semanales.',
    'Practicante TI: Luigi Armando Robles — desarrollo y soporte ERP.',
    'Apoderados/alumnos: beneficiarios; comunicación de cobro vía WhatsApp.',
])
H('Capacidades de negocio', 3)
bullet([
    'Formación marcial (clases, grados, valores).',
    'Gestión de alumnos (matrícula, ficha, plan, turno).',
    'Control operativo (horarios, asistencia P/A/J/R).',
    'Gestión financiera (mensualidades S/100, S/130, S/180).',
    'Gobierno y reportes (dashboard, comparativo mensual).',
])
H('Macroprocesos', 3)
mono('Consulta → Matrícula → Planificación turnos → Clase → Asistencia → Cuota mensual → Pago → Alertas (7 días) → Dashboard')
P('Figura 1: Arquitectura de negocio. Fuente: Elaboración propia.', bold=True)

# --- DIAGRAMA 2 ---
H('DIAGRAMA 2 — ARQUITECTURA DE APLICACIONES', 2)
P('Arquitectura JAMStack serverless: Next.js 16 (Vercel) + Supabase Cloud.')
H('Capas', 3)
bullet([
    'Cliente: navegador / PWA.',
    'Presentación: Next.js 16.2.4, React 19, Turbopack, UI iOS.',
    'Lógica: asistencia.service.js, pagoAlerts.service.js, auth, alumno, maestro.',
    'Integración: @supabase/supabase-js.',
    'Backend: PostgreSQL 17.6, PostgREST, Auth, RLS.',
    'Externos: WhatsApp, GitHub CI/CD.',
])
H('Módulos', 3)
bullet(['Login, Dashboard, Alumnos, Asistencia, Pagos, Maestros'])
H('Despliegue', 3)
mono('GitHub (main) → Vercel → Usuario HTTPS\n              ↓\n         Supabase Cloud')
P('Figura 2: Arquitectura de aplicaciones. Fuente: Elaboración propia.', bold=True)

# --- DIAGRAMA 3 ---
H('DIAGRAMA 3 — ARQUITECTURA DE DATOS', 2)
P('Datos centralizados en PostgreSQL 17 con dominios y RLS.')
H('Dominios', 3)
bullet([
    'Organización: sede, turno, maestro, maestro_turno.',
    'Académico: alumno, grado_marcial, historial_grados, clase, asistencia_alumno.',
    'Financiero: plan_mensualidad, pago, concepto_pago, metodo_pago.',
    'Seguridad: usuario, rol, RLS.',
])
H('Modelo ER (relaciones)', 3)
bullet([
    'sede (1)—turno (N)—clase (N); maestro (N)—turno (N) vía maestro_turno.',
    'alumno—turno, plan, grado; alumno—pago; alumno—asistencia—clase.',
    'pago.monto_final: columna GENERADA.',
])
H('Volúmenes v1.0', 3)
P('1 sede · 9 turnos · 60 alumnos · 366 clases · 770 asistencias · 144 pagos · 15 grados · 1 maestro 6° Dan.')
H('Flujo de datos', 3)
mono('UI → Reglas negocio → PostgreSQL → Reportes\nHTTPS → PostgREST → RLS → Tablas')
P('Figura 3: Arquitectura de datos. Fuente: Elaboración propia.', bold=True)

H('Nota', 2)
P('Diagramas gráficos Mermaid completos: ver DIAGRAMAS_ARQUITECTURA_ACCTKD.md. Exportar en https://mermaid.live como PNG para el informe.')

doc.save(OUT)
doc.save(OUT2)
print('Generado:', OUT)
print('Generado:', OUT2)
